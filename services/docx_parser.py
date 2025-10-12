import os
import tempfile
import mimetypes
import traceback
from typing import Tuple, Dict, List

# local imports
from services.pdf_parser import parse_pdf
# Note: avoid importing parse_docx from this module (self-import) as that
# creates a circular import when other modules try to import parse_docx.
# A lightweight helper `parse_docx` is provided at the end of this file.

# Optional import of textract for legacy formats
try:
    import textract  # type: ignore

    HAS_TEXTRACT = True
except Exception:
    HAS_TEXTRACT = False

# PyMuPDF for PDF image extraction
try:
    import fitz  # PyMuPDF

    HAS_FITZ = True
except Exception:
    HAS_FITZ = False

# python-docx for DOCX image extraction (already used in docx_parser)
try:
    import docx

    HAS_DOCCX = True
except Exception:
    HAS_DOCCX = False

# Pillow not required here anymore (OCR support removed)

# OCR: use the MistralOCR compatibility wrapper which now supports
# a Tesseract or Google Vision backend depending on env OCR_BACKEND.
try:
    from services.mistral_ocr import MistralOCR

    ocr_client = MistralOCR()
except Exception:
    ocr_client = None


def _save_bytes_to_tempfile(content_bytes: bytes, suffix: str = "") -> str:
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    with open(path, "wb") as f:
        f.write(content_bytes)
    return path


def _ocr_image_bytes(img_bytes: bytes) -> str:
    """
    Use MistralOCR to OCR a bytes object containing an image.
    Returns recognized text (may be empty string).
    """
    # Prefer the in-module ocr_client if available
    if ocr_client is None:
        return ""
    # write bytes to a temporary file and call extract_text which accepts a path
    try:
        tmp = _save_bytes_to_tempfile(img_bytes, suffix=".png")
        text = ocr_client.extract_text(tmp) or ""
        try:
            os.remove(tmp)
        except Exception:
            pass
        return text
    except Exception as e:
        print("OCR on image bytes failed:", e)
        return ""


def _extract_images_from_pdf(pdf_path: str) -> List[bytes]:
    """
    Extract images from a PDF and return list of image bytes.
    Requires PyMuPDF.
    """
    images = []
    if not HAS_FITZ:
        return images
    try:
        doc = fitz.open(pdf_path)
        for page_index in range(len(doc)):
            page = doc[page_index]
            image_list = page.get_images(full=True)
            for img in image_list:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                images.append(image_bytes)
        return images
    except Exception as e:
        print("PDF image extraction failed:", e)
        return images


def _extract_images_from_docx(docx_path: str) -> List[bytes]:
    """
    Extract images from a DOCX file using python-docx internals.
    Returns a list of image bytes.
    """
    images = []
    if not HAS_DOCCX:
        return images
    try:
        document = docx.Document(docx_path)
        # The document.part._rels contains relationships; images are stored with image parts
        # Iterate relationships and pick those with image content types
        rels = document.part._rels
        for rel in rels:
            rel_obj = rels[rel]
            if (
                "image" in rel_obj.target_ref.lower()
                or hasattr(rel_obj, "target_part")
                and getattr(rel_obj.target_part, "content_type", "").startswith(
                    "image/"
                )
            ):
                try:
                    img_part = rel_obj.target_part
                    img_bytes = img_part.blob
                    images.append(img_bytes)
                except Exception:
                    # fallback: try to read via rel_obj.target_ref relationship path
                    try:
                        # If target_ref is a path, try to open the file relative to docx package dir
                        pass
                    except Exception:
                        pass
        return images
    except Exception as e:
        print("DOCX image extraction failed:", e)
        return images


def _textract_fallback(file_path: str) -> str:
    """
    Use textract to extract text from many file formats (DOC, ODT, etc.).
    Requires textract to be installed on the system (and dependencies like antiword, catdoc sometimes).
    """
    if not HAS_TEXTRACT:
        raise RuntimeError(
            "textract is not installed — cannot process this file type as fallback."
        )
    try:
        text_bytes = textract.process(file_path)
        if isinstance(text_bytes, bytes):
            try:
                return text_bytes.decode("utf-8", errors="ignore")
            except Exception:
                return text_bytes.decode("latin-1", errors="ignore")
        return str(text_bytes)
    except Exception as e:
        print("textract fallback failed:", e)
        traceback.print_exc()
        return ""


def parse_document(file_path: str) -> Tuple[str, Dict]:
    """
    Main entrypoint to parse a document file into text.

    Returns:
      - combined_text (str): text extracted from document plus OCR of embedded images
      - manifest (dict): metadata about what was parsed (sources, counts, durations, errors)
    """
    manifest = {
        "file_path": file_path,
        "sources": [],
        "image_ocr_count": 0,
        "image_ocr_texts": [],
        "errors": [],
    }
    combined_text_chunks: List[str] = []

    # detect type based on extension and mimetype
    mime, _ = mimetypes.guess_type(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            # primary PDF text extraction (PyMuPDF parser)
            try:
                pdf_text = parse_pdf(file_path) or ""
                combined_text_chunks.append(pdf_text)
                manifest["sources"].append("pdf_text")
            except Exception as e:
                manifest["errors"].append(f"pdf_text_error:{e}")
            # extract images and OCR them
            try:
                images = _extract_images_from_pdf(file_path)
                for img_bytes in images:
                    ocr_text = _ocr_image_bytes(img_bytes)
                    if ocr_text:
                        combined_text_chunks.append(ocr_text)
                        manifest["image_ocr_texts"].append(ocr_text)
                manifest["image_ocr_count"] = len(images)
            except Exception as e:
                manifest["errors"].append(f"pdf_image_extract_error:{e}")

        elif ext in [".docx", ".docm", ".dotx"]:
            # Extract text with python-docx (already in docx_parser) - best for docx
            try:
                docx_text = parse_docx(file_path) or ""
                combined_text_chunks.append(docx_text)
                manifest["sources"].append("docx_text")
            except Exception as e:
                manifest["errors"].append(f"docx_text_error:{e}")

            # extract images and OCR them
            try:
                images = _extract_images_from_docx(file_path)
                for img_bytes in images:
                    ocr_text = _ocr_image_bytes(img_bytes)
                    if ocr_text:
                        combined_text_chunks.append(ocr_text)
                        manifest["image_ocr_texts"].append(ocr_text)
                manifest["image_ocr_count"] = len(images)
            except Exception as e:
                manifest["errors"].append(f"docx_image_extract_error:{e}")

        elif ext in [".odt", ".rtf", ".txt", ".md"]:
            # Try simple file read for text-like formats
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    txt = f.read()
                    combined_text_chunks.append(txt)
                    manifest["sources"].append("text_read")
            except Exception as e:
                manifest["errors"].append(f"text_read_error:{e}")
                # fallback to textract if available
                try:
                    txt = _textract_fallback(file_path)
                    combined_text_chunks.append(txt)
                    manifest["sources"].append("textract_fallback")
                except Exception as e2:
                    manifest["errors"].append(f"textract_error:{e2}")

        elif ext in [".doc"]:
            # Older .doc — use textract fallback (often requires antiword/catdoc on system)
            try:
                txt = _textract_fallback(file_path)
                combined_text_chunks.append(txt)
                manifest["sources"].append("textract_doc")
            except Exception as e:
                manifest["errors"].append(f"doc_textract_error:{e}")

        elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            # image file — run OCR directly
            try:
                with open(file_path, "rb") as f:
                    img_bytes = f.read()
                ocr_text = _ocr_image_bytes(img_bytes)
                combined_text_chunks.append(ocr_text)
                manifest["sources"].append("image_ocr")
                manifest["image_ocr_count"] = 1
                if ocr_text:
                    manifest["image_ocr_texts"].append(ocr_text)
            except Exception as e:
                manifest["errors"].append(f"image_ocr_error:{e}")

        else:
            # unknown extension — try text read, then textract fallback
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    txt = f.read()
                    combined_text_chunks.append(txt)
                    manifest["sources"].append("text_read_generic")
            except Exception:
                try:
                    txt = _textract_fallback(file_path)
                    combined_text_chunks.append(txt)
                    manifest["sources"].append("textract_generic")
                except Exception as e:
                    manifest["errors"].append(f"unknown_type_error:{e}")

    except Exception as e:
        manifest["errors"].append(
            f"parse_document_top_level_error:{e}\n{traceback.format_exc()}"
        )

    # Combine chunks and return
    combined_text = "\n\n".join(
        [c.strip() for c in combined_text_chunks if c and c.strip()]
    )

    return combined_text, manifest


if __name__ == "__main__":
    # quick local test (if running directly)
    print("Doc parser module test")
    test_path = input("Enter path to a sample doc/pdf/image: ").strip()
    if os.path.exists(test_path):
        text, meta = parse_document(test_path)
        print("---- Extracted text (first 1000 chars) ----")
        print(text[:1000])
        print("---- Manifest ----")
        import json

        print(json.dumps(meta, indent=2))
    else:
        print("File not found.")


def parse_docx(file_path: str) -> str:
    """
    Compatibility wrapper to extract text from a DOCX file.

    Kept as a top-level function so other modules can `from services.docx_parser
    import parse_docx` without triggering a circular import. It prefers python-docx
    extraction where available, and falls back to textract if installed.
    """
    # Try python-docx first (simple paragraph join)
    if HAS_DOCCX:
        try:
            document = docx.Document(file_path)
            paragraphs = [
                p.text for p in document.paragraphs if p.text and p.text.strip()
            ]
            return "\n\n".join(paragraphs)
        except Exception as e:
            print("python-docx extraction failed:", e)

    # Fallback to textract if available
    try:
        return _textract_fallback(file_path)
    except Exception:
        return ""
