"""
Utility functions for exporting resumes to PDF and DOCX formats
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    _HAS_REPORTLAB = True
except ImportError:
    _HAS_REPORTLAB = False
    print("Warning: reportlab not installed. PDF export will not be available.")
    print("Install with: pip install reportlab")


def export_to_docx(text_content, name="Resume", contact_info=None):
    """
    Export resume text to DOCX format
    
    Args:
        text_content: The resume text content
        name: Name for the document header
        contact_info: Optional dict with email, phone, location
    
    Returns:
        BytesIO object containing the DOCX file
    """
    try:
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add header with name
        if name and name != "Resume":
            header = doc.add_heading(name, level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.runs[0]
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(31, 78, 121)  # Professional blue
        
        # Add contact info if provided
        if contact_info:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_text = []
            if contact_info.get("email"):
                contact_text.append(contact_info["email"])
            if contact_info.get("phone"):
                contact_text.append(contact_info["phone"])
            if contact_info.get("location"):
                contact_text.append(contact_info["location"])
            
            run = contact_para.add_run(" | ".join(contact_text))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
            
            doc.add_paragraph()  # Spacer
        
        # Add main content
        # Split by sections and format
        lines = text_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header (all caps or title case with minimal words)
            if line.isupper() and len(line.split()) <= 4:
                # Section header
                heading = doc.add_heading(line.title(), level=2)
                run = heading.runs[0]
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(31, 78, 121)
            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                # Bullet point
                para = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                para.paragraph_format.left_indent = Inches(0.25)
                run = para.runs[0]
                run.font.size = Pt(11)
            else:
                # Regular paragraph
                para = doc.add_paragraph(line)
                run = para.runs[0]
                run.font.size = Pt(11)
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Error exporting to DOCX: {e}")
        # Return basic DOCX with just the text
        doc = Document()
        doc.add_paragraph(text_content)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


def export_to_pdf(text_content, name="Resume", contact_info=None):
    """
    Export resume text to PDF format using ReportLab
    
    Args:
        text_content: The resume text content
        name: Name for the document header
        contact_info: Optional dict with email, phone, location
    
    Returns:
        BytesIO object containing the PDF file
    """
    if not _HAS_REPORTLAB:
        raise RuntimeError("ReportLab not installed. Cannot export to PDF.")
    
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                              topMargin=0.5*inch, bottomMargin=0.5*inch,
                              leftMargin=0.75*inch, rightMargin=0.75*inch)
        
        # Container for the PDF elements
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(31/255, 78/255, 121/255),
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        contact_style = ParagraphStyle(
            'Contact',
            parent=styles['Normal'],
            fontSize=10,
            textColor=RGBColor(80/255, 80/255, 80/255),
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=RGBColor(31/255, 78/255, 121/255),
            spaceAfter=8,
            spaceBefore=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=4
        )
        
        # Add name header
        if name and name != "Resume":
            elements.append(Paragraph(name, title_style))
        
        # Add contact info
        if contact_info:
            contact_text = []
            if contact_info.get("email"):
                contact_text.append(contact_info["email"])
            if contact_info.get("phone"):
                contact_text.append(contact_info["phone"])
            if contact_info.get("location"):
                contact_text.append(contact_info["location"])
            
            elements.append(Paragraph(" | ".join(contact_text), contact_style))
        
        # Add main content
        lines = text_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 0.1*inch))
                continue
            
            # Escape HTML special characters for ReportLab
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Check if line is a section header
            if line.isupper() and len(line.split()) <= 4:
                elements.append(Paragraph(line.title(), heading_style))
            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                # Bullet point
                elements.append(Paragraph('• ' + line[1:].strip(), bullet_style))
            else:
                # Regular paragraph
                elements.append(Paragraph(line, body_style))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Error exporting to PDF: {e}")
        raise


def install_pdf_dependencies():
    """
    Helper function to install PDF export dependencies
    """
    return "pip install reportlab"
